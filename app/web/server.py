import asyncio
import base64
import json
import mimetypes
import os
import re
import subprocess
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal, Optional
from urllib.parse import quote, urlparse

from fastapi import (
    FastAPI,
    File,
    HTTPException,
    UploadFile,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from loguru import logger
from pydantic import BaseModel, Field

from app.agent.data_analysis import DataAnalysis
from app.agent.manus import Manus
from app.agent.task_control import (
    is_negative_research_conclusion,
    is_retrieval_objective,
)
from app.config import PROJECT_ROOT, config
from app.schema import Message
from app.team import ScopedManus, TeamCoordinator, TeamOutcome, TeamRole, TeamTask
from app.web.memory_rag import MemoryMatch, MemoryRecord, SemanticMemoryStore
from app.web.skill_matcher import SkillDocument, match_skills


STATIC_DIR = Path(__file__).parent / "static"
ASSETS_DIR = PROJECT_ROOT / "assets"
RUNS_FILE = config.workspace_root / "web-runs.json"
SKILLS_DIR = config.workspace_root / "skills"
UPLOADS_DIR = config.workspace_root / "uploads"
UPLOAD_METADATA_DIR = UPLOADS_DIR / ".metadata"
TERMINAL_STATUSES = {"completed", "error", "cancelled", "step_limit"}
DEFAULT_MAX_STEPS = 80
MAX_SKILL_CHARS = 12000
AUTO_SKILL_TOP_K = config.skill_config.top_k
AUTO_SKILL_MIN_SCORE = config.skill_config.min_score
MAX_UPLOAD_BYTES = 64 * 1024 * 1024
MAX_INLINE_IMAGE_BYTES = 8 * 1024 * 1024
MAX_UPLOADS_PER_RUN = 8
MAX_ATTACHMENT_PREVIEW_CHARS = 12000
ALLOWED_UPLOAD_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".png", ".jpg", ".jpeg"}
IMAGE_UPLOAD_EXTENSIONS = {".png", ".jpg", ".jpeg"}
DOCUMENT_UPLOAD_EXTENSIONS = {".docx", ".pdf", ".xlsx"}
EXACT_TEXT_PATTERNS = [
    re.compile(r"(?:评论|回复|发布|发送|输入|填写|私信)\s*[：:]\s*[\"“「『']([^\"”」』']{1,300})[\"”」』']"),
    re.compile(r"(?:评论|回复|发布|发送|输入|填写|私信)\s*[\"“「『']([^\"”」』']{1,300})[\"”」』']"),
    re.compile(r"(?:评论|回复|发布|发送|输入|填写|私信)\s*[：:]\s*([^\n]{1,300})"),
]
NO_FINAL_ANSWER_MESSAGE = (
    "任务已结束，但模型没有生成可展示的最终答案。运行过程已保留在右侧；" "请基于当前对话继续，让 MyManus 重新整理最终结论。"
)

app = FastAPI(title="MyManus Web", version="0.1.0")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
if ASSETS_DIR.exists():
    app.mount("/assets", StaticFiles(directory=ASSETS_DIR), name="assets")


def workspace_relative_path(value: str) -> Path:
    path = Path(value).expanduser()
    return path if path.is_absolute() else config.workspace_root / path


MEMORY_STORE = SemanticMemoryStore(
    storage_path=workspace_relative_path(config.memory_config.storage_path),
    cache_path=workspace_relative_path(config.memory_config.cache_path),
    provider=config.memory_config.provider,
    model_name=config.memory_config.model,
    model_cache_dir=workspace_relative_path(config.memory_config.model_cache_dir),
    base_url=config.memory_config.base_url,
    api_key=config.memory_config.api_key,
    api_key_env=config.memory_config.api_key_env,
    dimensions=config.memory_config.dimensions,
    max_records=config.memory_config.max_records,
    max_content_chars=config.memory_config.max_content_chars,
    fallback_to_sparse=config.memory_config.fallback_to_sparse,
)


def looks_like_raw_execution_log(value: Optional[str]) -> bool:
    if not value:
        return False
    text = value.lstrip()
    return bool(
        re.match(r"^Step\s+\d+:\s+", text)
        and (
            "Observed output of cmd `" in text
            or "### Ran Playwright code" in text
            or "Terminated: Reached max steps" in text
        )
    )


def display_answer(value: Optional[str], result: Optional[str] = None) -> Optional[str]:
    return append_artifact_links(value, result)


@app.on_event("startup")
async def load_saved_runs() -> None:
    load_runs()
    if config.memory_config.enabled:
        await asyncio.to_thread(backfill_memory_from_runs)


class RunRequest(BaseModel):
    prompt: str = Field(..., min_length=1)
    parent_run_id: Optional[str] = None
    skill_ids: list[str] = Field(default_factory=list)
    attachment_ids: list[str] = Field(default_factory=list)
    mode: Literal["single", "team"] = Field(default="single")


class SkillRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=80)
    content: str = Field(..., min_length=1, max_length=MAX_SKILL_CHARS)


@dataclass
class RunSession:
    id: str
    prompt: str
    max_steps: Optional[int]
    parent_run_id: Optional[str]
    skill_ids: list[str]
    auto_skill_matches: list[dict[str, Any]]
    attachments: list[dict[str, Any]]
    mode: str
    created_at: str
    updated_at: str
    memory_matches: list[dict[str, Any]] = field(default_factory=list)
    team: Optional[dict[str, Any]] = None
    status: str = "queued"
    result: Optional[str] = None
    answer: Optional[str] = None
    error: Optional[str] = None
    events: list[dict[str, Any]] = field(default_factory=list)
    subscribers: set[asyncio.Queue] = field(default_factory=set)
    task: Optional[asyncio.Task] = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "prompt": self.prompt,
            "max_steps": self.max_steps,
            "parent_run_id": self.parent_run_id,
            "skill_ids": self.skill_ids,
            "auto_skill_matches": self.auto_skill_matches,
            "memory_matches": self.memory_matches,
            "attachments": self.attachments,
            "mode": self.mode,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
            "team": self.team,
            "status": self.status,
            "result": self.result,
            "answer": display_answer(self.answer, self.result),
            "error": self.error,
            "event_count": len(self.events),
        }


runs: dict[str, RunSession] = {}
run_semaphore = asyncio.Semaphore(1)


def safe_upload_filename(value: str) -> str:
    name = Path(value or "attachment").name
    stem = Path(name).stem or "attachment"
    suffix = Path(name).suffix.lower()
    stem = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", stem).strip("._")
    if not stem:
        stem = "attachment"
    return f"{stem[:80]}{suffix}"


def format_bytes(size: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
        value /= 1024
    return f"{size} B"


def upload_metadata_path(upload_id: str) -> Path:
    return UPLOAD_METADATA_DIR / f"{upload_id}.json"


def read_upload_record(upload_id: str) -> Optional[dict[str, Any]]:
    if not re.fullmatch(r"[a-f0-9]{12}", upload_id or ""):
        return None

    metadata_path = upload_metadata_path(upload_id)
    if not metadata_path.is_file():
        return None

    try:
        record = json.loads(metadata_path.read_text(encoding="utf-8"))
    except Exception:
        return None

    path = Path(record.get("path") or "")
    if not path.is_absolute():
        path = config.workspace_root / path
    try:
        path.resolve().relative_to(config.workspace_root.resolve())
    except ValueError:
        return None
    if not path.is_file():
        return None

    record["absolute_path"] = str(path.resolve())
    record["url"] = workspace_file_url(str(path))
    return record


def selected_upload_records(upload_ids: list[str]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    seen: set[str] = set()
    for upload_id in upload_ids[:MAX_UPLOADS_PER_RUN]:
        record = read_upload_record(upload_id)
        if not record or record["id"] in seen:
            continue
        seen.add(record["id"])
        records.append(record)
    return records


def ensure_allowed_upload(path: Path, expected_suffix: str) -> None:
    suffix = expected_suffix.lower()
    if suffix not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Allowed: docx, pdf, xlsx, png, jpg, jpeg.",
        )

    try:
        if suffix == ".pdf":
            if not path.read_bytes()[:5].startswith(b"%PDF-"):
                raise ValueError("not a PDF file")
        elif suffix == ".docx":
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("not a valid DOCX file")
        elif suffix == ".xlsx":
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "xl/workbook.xml" not in names:
                raise ValueError("not a valid XLSX file")
        elif suffix in IMAGE_UPLOAD_EXTENSIONS:
            from PIL import Image

            with Image.open(path) as image:
                image.verify()
    except Exception as exc:
        path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Invalid {suffix} file.") from exc


async def save_uploaded_file(upload: UploadFile) -> dict[str, Any]:
    original_name = upload.filename or "attachment"
    safe_name = safe_upload_filename(original_name)
    suffix = Path(safe_name).suffix.lower()
    if suffix not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Allowed: docx, pdf, xlsx, png, jpg, jpeg.",
        )

    upload_id = uuid.uuid4().hex[:12]
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_METADATA_DIR.mkdir(parents=True, exist_ok=True)
    target = UPLOADS_DIR / f"{upload_id}-{safe_name}"

    size = 0
    try:
        with target.open("wb") as output:
            while True:
                chunk = await upload.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=413, detail="File is larger than 64 MB."
                    )
                output.write(chunk)
    finally:
        await upload.close()

    ensure_allowed_upload(target, suffix)
    content_type = (
        upload.content_type
        or mimetypes.guess_type(target.name)[0]
        or "application/octet-stream"
    )
    relative_path = (
        target.resolve().relative_to(config.workspace_root.resolve()).as_posix()
    )
    kind = "image" if suffix in IMAGE_UPLOAD_EXTENSIONS else suffix.lstrip(".")

    record = {
        "id": upload_id,
        "name": safe_name,
        "original_name": original_name,
        "extension": suffix.lstrip("."),
        "kind": kind,
        "content_type": content_type,
        "size": size,
        "size_label": format_bytes(size),
        "path": relative_path,
        "absolute_path": str(target.resolve()),
        "url": workspace_file_url(str(target)),
        "created_at": utc_now(),
    }
    upload_metadata_path(upload_id).write_text(
        json.dumps(record, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return record


def truncate_preview(value: str, limit: int) -> str:
    value = re.sub(r"\s+\n", "\n", value or "").strip()
    if len(value) <= limit:
        return value
    return f"{value[: limit // 2]}\n\n...[附件预览已截断]...\n\n{value[-limit // 2 :]}"


def docx_preview(path: Path, limit: int) -> str:
    from docx import Document

    document = Document(path)
    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    return truncate_preview("\n".join(parts), limit)


def xlsx_preview(path: Path, limit: int) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets[:5]:
        lines.append(f"Sheet: {sheet.title}")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value) for value in row]
            lines.append(" | ".join(values))
            if row_index >= 20:
                lines.append("... sheet preview truncated ...")
                break
    workbook.close()
    return truncate_preview("\n".join(lines), limit)


def pdf_preview(path: Path, limit: int) -> str:
    try:
        completed = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
    except Exception as exc:
        return f"PDF preview unavailable locally: {exc}"
    return truncate_preview(completed.stdout, limit)


def image_preview(path: Path) -> str:
    try:
        from PIL import Image

        with Image.open(path) as image:
            return f"Image metadata: format={image.format}, width={image.width}, height={image.height}, mode={image.mode}"
    except Exception as exc:
        return f"Image metadata unavailable: {exc}"


def stepfun_files_base_url() -> Optional[str]:
    llm_settings = config.llm.get("default")
    if not llm_settings:
        return None

    base_url = str(getattr(llm_settings, "base_url", "") or "")
    parsed = urlparse(base_url)
    if "stepfun" not in parsed.netloc:
        return None
    return f"{parsed.scheme or 'https'}://{parsed.netloc}/v1"


def stepfun_api_key() -> Optional[str]:
    llm_settings = config.llm.get("default")
    if not llm_settings:
        return None
    api_key = str(getattr(llm_settings, "api_key", "") or "")
    return api_key or None


def delete_stepfun_file(file_id: str) -> None:
    base_url = stepfun_files_base_url()
    api_key = stepfun_api_key()
    if not base_url or not api_key or not file_id:
        return

    try:
        import requests

        requests.delete(
            f"{base_url}/files/{file_id}",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=10,
        )
    except Exception:
        return


def stepfun_file_extract_preview(record: dict[str, Any], limit: int) -> str:
    base_url = stepfun_files_base_url()
    api_key = stepfun_api_key()
    if not base_url or not api_key:
        return ""

    path = Path(record["absolute_path"])
    if path.suffix.lower() not in DOCUMENT_UPLOAD_EXTENSIONS:
        return ""

    file_id = ""
    try:
        import time

        import requests

        with path.open("rb") as handle:
            upload_response = requests.post(
                f"{base_url}/files",
                headers={"Authorization": f"Bearer {api_key}"},
                data={"purpose": "file-extract"},
                files={"file": (record["name"], handle)},
                timeout=45,
            )
        upload_response.raise_for_status()
        upload_payload = upload_response.json()
        file_id = str(upload_payload.get("id") or "")
        if not file_id:
            return ""

        status = str(upload_payload.get("status") or "")
        for _ in range(12):
            if status == "success":
                break
            time.sleep(1)
            retrieve_response = requests.get(
                f"{base_url}/files/{file_id}",
                headers={"Authorization": f"Bearer {api_key}"},
                timeout=20,
            )
            retrieve_response.raise_for_status()
            status = str(retrieve_response.json().get("status") or "")

        if status != "success":
            return ""

        content_response = requests.get(
            f"{base_url}/files/{file_id}/content",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=45,
        )
        content_response.raise_for_status()
        content = content_response.text.strip()
        if not content:
            return ""
        return truncate_preview(content, limit)
    except Exception as exc:
        logger.info(f"StepFun file extraction fallback for {path.name}: {exc}")
        return ""
    finally:
        delete_stepfun_file(file_id)


def attachment_preview(record: dict[str, Any], limit: int) -> str:
    path = Path(record["absolute_path"])
    suffix = path.suffix.lower()
    try:
        if suffix in DOCUMENT_UPLOAD_EXTENSIONS:
            official_preview = stepfun_file_extract_preview(record, limit)
            if official_preview:
                return official_preview
        if suffix == ".docx":
            return docx_preview(path, limit)
        if suffix == ".xlsx":
            return xlsx_preview(path, limit)
        if suffix == ".pdf":
            return pdf_preview(path, limit)
        if suffix in IMAGE_UPLOAD_EXTENSIONS:
            return image_preview(path)
    except Exception as exc:
        return f"Preview unavailable: {exc}"
    return ""


def gmail_auth_status() -> dict[str, Any]:
    gmail_home = Path.home() / ".gmail-mcp"
    return {
        "configured": "gmail" in config.mcp_config.servers,
        "oauth_keys": (gmail_home / "gcp-oauth.keys.json").is_file(),
        "credentials": (gmail_home / "credentials.json").is_file(),
        "credentials_dir": str(gmail_home),
        "auth_command": "npm run gmail:auth",
    }


def attachments_prompt(records: list[dict[str, Any]]) -> str:
    if not records:
        return ""

    remaining_preview = MAX_ATTACHMENT_PREVIEW_CHARS
    lines = [
        "用户本轮上传了以下附件。请把这些文件当作任务输入来使用；不要要求用户重新上传。",
        "文档附件会优先尝试阶跃星辰 Files API 的 file-extract 获取纯文本预览，失败时回退本地解析。",
        "需要深入分析时，可以使用 Python 工具按 absolute_path 读取文件：docx 用 python-docx，xlsx 用 openpyxl，pdf 可优先尝试 pdftotext，图片可用 Pillow 获取尺寸；图片附件也会在模型支持时作为视觉输入提供。",
    ]

    for index, record in enumerate(records, start=1):
        lines.extend(
            [
                "",
                f"附件 {index}:",
                f"- id: {record['id']}",
                f"- name: {record['name']}",
                f"- type: {record['extension']}",
                f"- size: {record.get('size_label') or format_bytes(int(record.get('size') or 0))}",
                f"- absolute_path: {record['absolute_path']}",
                f"- download_url: {record['url']}",
            ]
        )
        if remaining_preview > 0:
            preview = attachment_preview(record, min(remaining_preview, 4000))
            if preview:
                remaining_preview -= len(preview)
                lines.extend(["- preview:", "```", preview, "```"])

    return "\n".join(lines)


def image_records_for_direct_input(
    records: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    result = []
    for record in records:
        if record.get("kind") != "image":
            continue
        if int(record.get("size") or 0) > MAX_INLINE_IMAGE_BYTES:
            continue
        result.append(record)
    return result[:4]


def add_uploaded_images_to_agent(agent: Manus, records: list[dict[str, Any]]) -> None:
    for record in image_records_for_direct_input(records):
        path = Path(record["absolute_path"])
        try:
            encoded = base64.b64encode(path.read_bytes()).decode("utf-8")
        except Exception as exc:
            logger.warning(f"Failed to attach uploaded image {path}: {exc}")
            continue

        agent.memory.add_message(
            Message.user_message(
                (
                    "上传图片附件，可直接视觉理解。"
                    f"\nname: {record['name']}"
                    f"\nabsolute_path: {record['absolute_path']}"
                ),
                base64_image=encoded,
                image_mime_type=record.get("content_type")
                or mimetypes.guess_type(path.name)[0]
                or "image/jpeg",
            )
        )


def run_to_storage(run: RunSession) -> dict[str, Any]:
    data = run.to_dict()
    data.update(
        {
            "result": run.result,
            "answer": display_answer(run.answer, run.result),
            "error": run.error,
        }
    )
    return data


def save_runs() -> None:
    try:
        RUNS_FILE.parent.mkdir(parents=True, exist_ok=True)
        stored = [
            run_to_storage(run)
            for run in sorted(
                runs.values(), key=lambda item: item.created_at, reverse=True
            )
            if run.status in TERMINAL_STATUSES
        ][:50]
        RUNS_FILE.write_text(
            json.dumps(stored, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as exc:
        logger.warning(f"Failed to save web runs: {exc}")


def load_runs() -> None:
    if not RUNS_FILE.exists():
        return

    try:
        stored = json.loads(RUNS_FILE.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning(f"Failed to load web runs: {exc}")
        return

    for item in stored:
        try:
            run = RunSession(
                id=item["id"],
                prompt=item["prompt"],
                max_steps=item.get("max_steps"),
                parent_run_id=item.get("parent_run_id"),
                skill_ids=item.get("skill_ids") or [],
                auto_skill_matches=item.get("auto_skill_matches") or [],
                attachments=item.get("attachments") or [],
                mode=item.get("mode")
                if item.get("mode") in {"single", "team"}
                else "single",
                created_at=item["created_at"],
                updated_at=item["updated_at"],
                memory_matches=item.get("memory_matches") or [],
                team=item.get("team"),
                status=item.get("status", "completed"),
                result=item.get("result"),
                answer=item.get("answer"),
                error=item.get("error"),
            )
            runs[run.id] = run
        except Exception as exc:
            logger.warning(f"Skipped invalid saved web run: {exc}")


def root_run_id(run: RunSession) -> str:
    current = run
    seen: set[str] = set()
    while current.parent_run_id and current.parent_run_id in runs:
        if current.id in seen or current.parent_run_id in seen:
            break
        seen.add(current.id)
        current = runs[current.parent_run_id]
    return current.id


def thread_run_ids(run_id: str) -> list[str]:
    run = runs.get(run_id)
    if not run:
        return []

    root_id = root_run_id(run)
    return [
        candidate.id for candidate in runs.values() if root_run_id(candidate) == root_id
    ]


def conversation_runs(run_id: str) -> list[RunSession]:
    ids = set(thread_run_ids(run_id))
    return sorted(
        (run for run in runs.values() if run.id in ids),
        key=lambda item: (item.created_at, item.updated_at),
    )


def memory_record_from_run(run: RunSession) -> Optional[MemoryRecord]:
    if run.status not in {"completed", "step_limit"}:
        return None
    answer = str(display_answer(run.answer, run.result) or "").strip()
    if not answer and not run.result:
        return None
    observations = str(run.result or "")
    web_research_evidence = bool(
        re.search(
            r"cmd `(?:mcp_stepsearch_|mcp_playwright_browser_|web_search|web_fetch|crawl)",
            observations,
            flags=re.IGNORECASE,
        )
    )
    negative_research = (
        is_retrieval_objective(run.prompt)
        and web_research_evidence
        and is_negative_research_conclusion(answer)
    )
    direct_source_read = bool(
        re.search(
            r"cmd `(?:mcp_stepsearch_web_fetch|mcp_playwright_browser_"
            r"(?:snapshot|evaluate))`",
            observations,
            flags=re.IGNORECASE,
        )
    )
    retrieval_eligible = not (negative_research and not direct_source_read)
    quality = "completed" if retrieval_eligible else "unverified_negative"
    return MemoryRecord(
        id=run.id,
        conversation_id=root_run_id(run),
        run_id=run.id,
        created_at=run.created_at,
        task=truncate_context(run.prompt, 4000),
        answer=truncate_context(answer, 8000),
        observations=truncate_context(observations, 6000),
        status=run.status,
        quality=quality,
        retrieval_eligible=retrieval_eligible,
    )


def remember_run(run: RunSession) -> None:
    if not config.memory_config.enabled:
        return
    record = memory_record_from_run(run)
    if record:
        MEMORY_STORE.upsert(record)


def backfill_memory_from_runs() -> None:
    for run in sorted(runs.values(), key=lambda item: item.created_at):
        remember_run(run)


def serialize_memory_match(match: MemoryMatch) -> dict[str, Any]:
    record = match.record
    return {
        "id": record.id,
        "conversation_id": record.conversation_id,
        "run_id": record.run_id,
        "created_at": record.created_at,
        "task": record.task,
        "answer": record.answer,
        "observations": record.observations,
        "status": record.status,
        "quality": record.quality,
        "score": match.score,
        "retrieval_method": match.retrieval_method,
        "embedding_model": match.embedding_model,
    }


async def retrieve_relevant_memories(
    prompt: str, *, exclude_run_ids: set[str] | None = None
) -> list[dict[str, Any]]:
    if not config.memory_config.enabled:
        return []
    matches = await asyncio.to_thread(
        MEMORY_STORE.search,
        prompt,
        exclude_run_ids=exclude_run_ids or set(),
        top_k=config.memory_config.top_k,
        min_score=config.memory_config.min_score,
    )
    return [serialize_memory_match(match) for match in matches]


def recent_thread_runs() -> list[RunSession]:
    latest_by_root: dict[str, RunSession] = {}
    for run in runs.values():
        root_id = root_run_id(run)
        current = latest_by_root.get(root_id)
        if current is None or (run.updated_at, run.created_at) > (
            current.updated_at,
            current.created_at,
        ):
            latest_by_root[root_id] = run

    return sorted(
        latest_by_root.values(),
        key=lambda item: (item.updated_at, item.created_at),
        reverse=True,
    )


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def secret_values() -> list[str]:
    values: list[str] = []
    for llm_settings in config.llm.values():
        api_key = getattr(llm_settings, "api_key", None)
        if api_key:
            values.append(api_key)

    for server in config.mcp_config.servers.values():
        if server.env:
            values.extend(value for value in server.env.values() if value)

    return sorted({value for value in values if len(value) >= 8}, key=len, reverse=True)


def redact(value: Any) -> Any:
    if isinstance(value, list):
        return [redact(item) for item in value]
    if isinstance(value, dict):
        return {key: redact(item) for key, item in value.items()}
    if not isinstance(value, str):
        return value

    redacted = value
    for secret in secret_values():
        redacted = redacted.replace(secret, "<redacted>")

    redacted = re.sub(
        r"(PLAYWRIGHT_MCP_EXTENSION_TOKEN=)[A-Za-z0-9_.~+-]+",
        r"\1<redacted>",
        redacted,
    )
    redacted = re.sub(r"([?&]token=)[^)&\s]+", r"\1<redacted>", redacted)
    redacted = re.sub(
        r"(api[_-]?key['\"]?\s*[:=]\s*['\"]?)[^'\"\s,]+",
        r"\1<redacted>",
        redacted,
        flags=re.IGNORECASE,
    )
    return redacted


def redact_payload(payload: dict[str, Any]) -> dict[str, Any]:
    return {key: redact(value) for key, value in payload.items()}


def normalize_skill_id(value: str) -> str:
    raw = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", value.strip(), flags=re.UNICODE)
    raw = raw.strip("-_.")
    if not raw:
        raw = f"skill-{uuid.uuid4().hex[:8]}"
    return raw[:64]


def skill_markdown_path(skill_id: str) -> Path:
    normalized = normalize_skill_id(skill_id)
    target = (SKILLS_DIR / normalized / "SKILL.md").resolve()
    root = SKILLS_DIR.resolve()
    try:
        target.parent.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="Skill not found.") from exc
    return target


def parse_skill_metadata(skill_id: str, content: str) -> dict[str, Any]:
    lines = [line.strip() for line in content.splitlines()]
    heading = next((line[2:].strip() for line in lines if line.startswith("# ")), "")
    name = heading or skill_id
    summary_lines: list[str] = []
    seen_heading = False
    for line in lines:
        if line.startswith("#"):
            seen_heading = True
            continue
        if not line:
            if summary_lines:
                break
            continue
        if seen_heading or not heading:
            summary_lines.append(line)
        if len(" ".join(summary_lines)) >= 180:
            break
    summary = re.sub(r"\s+", " ", " ".join(summary_lines)).strip()
    return {
        "id": skill_id,
        "name": name[:80],
        "summary": summary[:220],
        "path": str(skill_markdown_path(skill_id)),
    }


def list_skill_records() -> list[dict[str, Any]]:
    SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    records: list[dict[str, Any]] = []
    for path in sorted(SKILLS_DIR.glob("*/SKILL.md")):
        try:
            skill_id = path.parent.name
            content = path.read_text(encoding="utf-8")
            records.append(parse_skill_metadata(skill_id, content))
        except OSError:
            continue
    return records


def list_skill_documents() -> list[SkillDocument]:
    SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    documents: list[SkillDocument] = []
    for path in sorted(SKILLS_DIR.glob("*/SKILL.md")):
        try:
            skill_id = path.parent.name
            content = path.read_text(encoding="utf-8")
            metadata = parse_skill_metadata(skill_id, content)
            documents.append(
                SkillDocument(
                    id=skill_id,
                    name=metadata["name"],
                    summary=metadata["summary"],
                    content=content,
                )
            )
        except OSError:
            continue
    return documents


def read_skill_content(skill_id: str) -> str:
    path = skill_markdown_path(skill_id)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Skill not found.")
    return path.read_text(encoding="utf-8")


def valid_skill_ids(values: list[str]) -> list[str]:
    seen: set[str] = set()
    valid: list[str] = []
    for value in values or []:
        skill_id = normalize_skill_id(value)
        if skill_id in seen:
            continue
        if skill_markdown_path(skill_id).is_file():
            seen.add(skill_id)
            valid.append(skill_id)
    return valid[:8]


def serialize_skill_match(match) -> dict[str, Any]:
    return {
        "id": match.id,
        "name": match.name,
        "summary": match.summary,
        "score": match.score,
        "retrieval_method": match.retrieval_method,
        "matched_terms": list(match.matched_terms),
    }


def skill_match_query(prompt: str, attachments: list[dict[str, Any]]) -> str:
    parts = [prompt]
    for record in attachments or []:
        parts.append(
            " ".join(
                str(value)
                for value in [
                    record.get("name"),
                    record.get("extension"),
                    record.get("content_type"),
                ]
                if value
            )
        )
    return "\n".join(part for part in parts if part)


async def retrieve_auto_skill_matches(
    prompt: str,
    selected_skill_ids: list[str],
    attachments: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    matches = await asyncio.to_thread(
        match_skills,
        skill_match_query(prompt, attachments),
        list_skill_documents(),
        exclude_ids=set(selected_skill_ids),
        top_k=AUTO_SKILL_TOP_K,
        min_score=AUTO_SKILL_MIN_SCORE,
    )
    return [serialize_skill_match(match) for match in matches]


def combined_skill_ids(run: RunSession) -> list[str]:
    return valid_skill_ids(
        [
            *run.skill_ids,
            *(str(match.get("id") or "") for match in run.auto_skill_matches),
        ]
    )


def selected_skills_prompt(
    skill_ids: list[str],
    auto_skill_matches: list[dict[str, Any]] | None = None,
) -> str:
    valid_ids = valid_skill_ids(skill_ids)
    if not valid_ids:
        return ""

    auto_by_id = {
        str(match.get("id")): match
        for match in auto_skill_matches or []
        if match.get("id")
    }
    parts = [
        "以下是本轮任务开始前已加载的自定义 Skills，包含用户手动选择和基于关键词匹配自动加载的 Skills。"
        "请先阅读并遵守这些 SKILL.md，但如果它们与用户本轮明确要求冲突，以用户本轮要求为准。",
    ]
    for index, skill_id in enumerate(valid_ids, start=1):
        content = truncate_context(read_skill_content(skill_id), MAX_SKILL_CHARS)
        metadata = parse_skill_metadata(skill_id, content)
        auto_match = auto_by_id.get(skill_id)
        if auto_match:
            source = (
                f"TF-IDF 关键词匹配，score={auto_match.get('score')}, "
                f"matched_terms={', '.join(auto_match.get('matched_terms') or []) or '-'}"
            )
        else:
            source = "用户手动选择"
        parts.extend(
            [
                "",
                f"Skill {index}: {metadata['name']} ({skill_id})",
                f"Source: {source}",
                "```markdown",
                content,
                "```",
            ]
        )
    return "\n".join(parts)


def publish(run: RunSession, event_type: str, **payload: Any) -> None:
    event = {
        "id": len(run.events) + 1,
        "time": utc_now(),
        "type": event_type,
        **redact_payload(payload),
    }
    run.updated_at = event["time"]
    run.events.append(event)
    if len(run.events) > 1000:
        run.events = run.events[-1000:]

    for queue in list(run.subscribers):
        try:
            queue.put_nowait(event)
        except asyncio.QueueFull:
            run.subscribers.discard(queue)


def set_status(run: RunSession, status: str, **payload: Any) -> None:
    run.status = status
    publish(run, "status", status=status, **payload)
    if status in TERMINAL_STATUSES:
        save_runs()


def prune_runs() -> None:
    if len(runs) <= 50:
        return
    finished = [
        run
        for run in sorted(runs.values(), key=lambda item: item.created_at)
        if run.status in TERMINAL_STATUSES
    ]
    for run in finished[: len(runs) - 50]:
        runs.pop(run.id, None)


def extract_json_from_text(value: str) -> Optional[dict[str, Any]]:
    start = value.find("{")
    end = value.rfind("}")
    if start < 0 or end <= start:
        return None
    try:
        parsed = json.loads(value[start : end + 1])
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def extract_plain_step_answer(result: str) -> str:
    """Return the last non-tool natural-language step from agent.run output."""
    matches = list(
        re.finditer(r"(?:^|\n)Step\s+\d+:\s*([\s\S]*?)(?=\nStep\s+\d+:|\Z)", result)
    )
    for match in reversed(matches):
        content = match.group(1).strip()
        if not content:
            continue
        if content.startswith("Observed output of cmd `"):
            continue
        if content.startswith("Error executing"):
            continue
        if (
            "The interaction has been completed with status:" in content
            and len(content) < 200
        ):
            continue
        return content
    return ""


def workspace_file_url(path: str) -> str:
    root = config.workspace_root.resolve()
    raw = Path(path).expanduser()
    target = raw if raw.is_absolute() else root / raw
    try:
        relative = target.resolve().relative_to(root)
    except ValueError:
        return ""
    return f"/api/files/{quote(relative.as_posix(), safe='/')}"


def extract_artifact_answer(result: str) -> str:
    artifacts: list[dict[str, Any]] = []
    pattern = r"Step\s+\d+:\s+Observed output of cmd `([^`]+)` executed:\n([\s\S]*?)(?=\nStep\s+\d+:|\Z)"
    for match in re.finditer(pattern, result):
        tool_name = match.group(1)
        if tool_name not in {"create_word_document", "create_excel_workbook"}:
            continue
        payload = extract_json_from_text(match.group(2))
        if not payload or not payload.get("path"):
            continue
        artifacts.append(payload)

    # Executors may create a requested deliverable through Python or another
    # implementation tool. Ground those claims in files that actually exist inside
    # the workspace instead of relying on the worker's prose or a specific tool name.
    known_paths = {str(item.get("path") or "") for item in artifacts}
    workspace = config.workspace_root.resolve()
    path_pattern = re.compile(
        rf"(?:{re.escape(str(workspace))}/|workspace/)"
        r"[^\n\r\"'`]+?\.(?:docx|xlsx|pdf)",
        re.IGNORECASE,
    )
    for match in path_pattern.finditer(result or ""):
        raw_path = match.group(0).strip()
        raw = Path(raw_path)
        if raw.is_absolute():
            candidate = raw
        elif raw.parts and raw.parts[0] == workspace.name:
            candidate = workspace.joinpath(*raw.parts[1:])
        else:
            candidate = workspace / raw
        try:
            resolved = candidate.resolve()
            resolved.relative_to(workspace)
        except (OSError, ValueError):
            continue
        if not resolved.is_file() or str(resolved) in known_paths:
            continue
        artifact_type = {
            ".docx": "word_document",
            ".xlsx": "excel_workbook",
            ".pdf": "pdf_document",
        }.get(resolved.suffix.lower(), "file")
        artifacts.append({"type": artifact_type, "path": str(resolved)})
        known_paths.add(str(resolved))

    if not artifacts:
        return ""

    lines = ["任务已完成，并生成了以下产物："]
    for item in artifacts:
        artifact_type = {
            "excel_workbook": "Excel",
            "word_document": "Word",
            "pdf_document": "PDF",
        }.get(str(item.get("type") or ""), str(item.get("type") or "文件"))
        file_path = str(item["path"])
        file_url = workspace_file_url(file_path)
        file_name = Path(file_path).name
        if file_url:
            lines.append(f"- {artifact_type}: [{file_name}]({file_url})")
            lines.append(f"  路径：`{file_path}`")
        else:
            lines.append(f"- {artifact_type}: `{file_path}`")
        if item.get("sheets"):
            lines.append(f"  工作表：{', '.join(str(sheet) for sheet in item['sheets'])}")
        if item.get("data_rows") is not None:
            lines.append(f"  数据行数：{item['data_rows']}")
    return "\n".join(lines)


def append_artifact_links(
    answer: Optional[str], result: Optional[str]
) -> Optional[str]:
    if not answer or not result:
        return answer

    artifact_answer = extract_artifact_answer(result)
    if not artifact_answer:
        return answer

    if "/api/files/" in answer:
        return answer

    artifact_urls = re.findall(r"\]\((/api/files/[^)]+)\)", artifact_answer)
    if any(url in answer for url in artifact_urls):
        return answer

    return f"{answer.rstrip()}\n\n---\n{artifact_answer}"


def remove_private_thinking(value: str) -> str:
    return re.sub(r"(?is)<think>.*?</think>", "", value)


def extract_user_answer_content(content: Optional[str]) -> str:
    if not content:
        return ""

    text = remove_private_thinking(content).strip()
    if not text:
        return ""

    thought_match = re.match(
        r"(?is)^\s*(?:Thought|思考|思考摘要|思路摘要|计划)\s*[:：]\s*(.*)$",
        text,
    )
    if thought_match:
        rest = thought_match.group(1).strip()
        paragraph_stop = re.search(r"\n\s*\n", rest)
        if not paragraph_stop:
            return ""
        text = rest[paragraph_stop.end() :].strip()

    text = re.sub(
        r"(?im)^\s*(?:Final Answer|最终答案|答案|回复)\s*[:：]\s*",
        "",
        text,
    ).strip()
    if text.startswith("Observed output of cmd `"):
        return ""
    if "The interaction has been completed with status:" in text and len(text) < 200:
        return ""
    return text


def extract_final_answer(agent: Manus, fallback: str) -> str:
    """Return a real user-facing final answer instead of tool-call preambles."""
    artifact_answer = extract_artifact_answer(fallback)
    structured_final_answer = extract_user_answer_content(
        getattr(agent, "final_answer", None)
    )
    if structured_final_answer:
        return (
            append_artifact_links(structured_final_answer, fallback)
            or structured_final_answer
        )

    for message in reversed(agent.memory.messages):
        if message.role != "assistant" or not message.content:
            continue
        if message.tool_calls and not any(
            call.function.name == "terminate" for call in message.tool_calls
        ):
            continue
        content = extract_user_answer_content(message.content)
        if content:
            return append_artifact_links(content, fallback) or content

    plain_step_answer = extract_plain_step_answer(fallback)
    if plain_step_answer:
        return append_artifact_links(plain_step_answer, fallback) or plain_step_answer

    if artifact_answer:
        return artifact_answer

    return NO_FINAL_ANSWER_MESSAGE


def truncate_context(value: Optional[str], limit: int) -> str:
    if not value:
        return ""
    text = value.strip()
    if len(text) <= limit:
        return text
    return f"{text[: limit // 2]}\n\n...[中间内容已截断]...\n\n{text[-limit // 2 :]}"


def extract_required_exact_text(prompt: str) -> str:
    for pattern in EXACT_TEXT_PATTERNS:
        match = pattern.search(prompt or "")
        if match:
            return match.group(1).strip().strip("\"“”「」『』'")
    return ""


def append_exact_text_constraints(prompt: str, source_prompt: str) -> str:
    exact_text = extract_required_exact_text(source_prompt)
    if not exact_text:
        return prompt
    return "\n\n---\n\n".join(
        [
            prompt,
            (
                "硬性执行约束：\n"
                f"- 用户明确指定的原文是：{exact_text}\n"
                "- 如果本轮任务要求输入、填写、评论、回复、发布、发送或私信，必须逐字逐标点使用这段原文。\n"
                "- 禁止改写、润色、同义替换、补充解释、修正错别字或换成自认为更合适的内容。\n"
                "- 完成前必须用工具证据验证外部页面或结果中出现的文本与这段原文完全一致；不一致时不要 terminate(success)。"
            ),
        ]
    )


def relevant_memories_prompt(matches: list[dict[str, Any]]) -> str:
    if not matches:
        return ""
    parts = [
        "以下内容由 Agent Memory RAG 从历史任务中按语义相关性召回。",
        "它们只用于继承相关事实、用户偏好、历史产物和已验证经验；历史内容可能已经过期，若与本轮要求或当前工具状态冲突，以本轮要求和当前状态为准。",
    ]
    for index, match in enumerate(matches, start=1):
        parts.extend(
            [
                "",
                f"Memory {index}: run={match.get('run_id')} score={match.get('score')} model={match.get('embedding_model') or '-'}",
                f"历史任务：{truncate_context(match.get('task'), 2500)}",
                f"历史结果：{truncate_context(match.get('answer'), 4500)}",
            ]
        )
        observations = truncate_context(match.get("observations"), 2500)
        if observations:
            parts.append(f"关键执行记录：{observations}")
    return "\n".join(parts)


def attach_run_context(run: RunSession, prompt: str) -> str:
    skills_prompt = selected_skills_prompt(
        combined_skill_ids(run), run.auto_skill_matches
    )
    memories_prompt = relevant_memories_prompt(run.memory_matches)
    attachment_prompt = attachments_prompt(run.attachments)
    sections = []
    if memories_prompt:
        sections.append(memories_prompt)
    if skills_prompt:
        sections.append(skills_prompt)
    if attachment_prompt:
        sections.append(attachment_prompt)
    if sections:
        sections.append(f"本轮用户任务：\n{prompt}")
        return "\n\n---\n\n".join(sections)
    return prompt


def build_execution_prompt(run: RunSession) -> str:
    if not run.parent_run_id:
        return attach_run_context(
            run, append_exact_text_constraints(run.prompt, run.prompt)
        )

    if run.parent_run_id not in runs:
        return attach_run_context(
            run, append_exact_text_constraints(run.prompt, run.prompt)
        )

    history = [item for item in conversation_runs(run.id) if item.id != run.id]
    if not history:
        return attach_run_context(
            run, append_exact_text_constraints(run.prompt, run.prompt)
        )

    parts = [
        "你正在继续当前 MyManus Web 对话。",
        "请把下面的历史对话当作已经发生的工作基础，不要无意义地从零开始。",
        "如果浏览器仍停留在上一轮相关页面，可以优先基于当前页面继续；如果页面状态丢失，再自行恢复必要页面。",
        "本轮用户追加要求是最高优先级；历史对话只用于解析指代、继承已确认事实和延续未完成工作。",
        "如果本轮出现“他们”“上述”“里面”“这两个”“这些”等指代，必须先根据最近历史解析成具体对象，再围绕这些对象完成本轮问题。",
        "如果无法从历史中唯一确定指代对象，应该要求用户澄清；不要把任务改写成背景介绍、主页查找或无关总结。",
        "最终答案必须直接回答本轮问题。找到网页、主页、资料源或中间证据都只是过程，不等于任务完成。",
        "",
        "历史对话：",
    ]

    for index, item in enumerate(history[-8:], start=max(1, len(history) - 7)):
        answer = truncate_context(item.answer or item.error, 8000)
        result = truncate_context(item.result, 4000)
        parts.extend(
            [
                "",
                f"第 {index} 轮任务 ID：{item.id}",
                f"第 {index} 轮状态：{item.status}",
                f"第 {index} 轮用户任务：",
                item.prompt,
            ]
        )
        if answer:
            parts.extend(["", f"第 {index} 轮最终答案：", answer])
        if result and item.status == "step_limit":
            parts.extend(["", f"第 {index} 轮执行摘要（因步数耗尽，仅保留截断上下文）：", result])
        elif result and display_answer(item.answer) == NO_FINAL_ANSWER_MESSAGE:
            parts.extend(
                [
                    "",
                    f"第 {index} 轮执行摘要（模型未生成最终答案，仅保留截断上下文）：",
                    truncate_context(item.result, 12000),
                ]
            )

    if len(history) > 8:
        parts.extend(
            [
                "",
                f"注：共有 {len(history)} 轮历史，上面保留最近 8 轮作为可执行上下文。",
            ]
        )

    parts.extend(["", "本轮用户追加要求：", run.prompt])
    return attach_run_context(
        run, append_exact_text_constraints("\n".join(parts), run.prompt)
    )


async def execute_team_run(run: RunSession, execution_prompt: str) -> TeamOutcome:
    async def worker_factory(role: TeamRole, task: TeamTask):
        if role == TeamRole.DATA:
            worker = DataAnalysis()
        else:
            worker = await ScopedManus.create_for_role(role.value)

        add_uploaded_images_to_agent(worker, run.attachments)
        tool_names = sorted(worker.available_tools.tool_map)
        publish(
            run,
            "team_agent",
            task_id=task.id,
            role=role.value,
            agent=worker.name,
            tool_count=len(tool_names),
            tools=tool_names,
        )
        return worker

    def team_event_sink(event_type: str, payload: dict[str, Any]) -> None:
        publish(run, event_type, **payload)

    max_steps = run.max_steps or DEFAULT_MAX_STEPS
    coordinator = TeamCoordinator(
        worker_factory=worker_factory,
        event_sink=team_event_sink,
        worker_max_steps=max(8, min(24, max_steps // 2)),
        worker_step_limits={
            TeamRole.BROWSER: max(12, min(32, max_steps // 2)),
            TeamRole.DATA: max(8, min(20, max_steps // 3)),
            TeamRole.GENERAL: max(10, min(24, max_steps // 3)),
        },
        cancel_requested=lambda: run.status == "cancelling",
    )
    outcome = await coordinator.execute(run.prompt, execution_prompt)
    run.team = redact(outcome.snapshot)
    run.result = redact(outcome.trace)
    run.answer = redact(
        append_artifact_links(outcome.answer, outcome.trace) or outcome.answer
    )
    publish(run, "answer", content=run.answer)
    publish(run, "result", content=run.result)
    return outcome


async def execute_run(run: RunSession) -> None:
    agent: Optional[Manus] = None
    sink_id: Optional[int] = None

    def log_sink(message) -> None:
        publish(run, "log", level=message.record["level"].name, message=str(message))

    try:
        if run.status == "cancelling":
            raise asyncio.CancelledError
        set_status(run, "waiting")
        async with run_semaphore:
            if run.status == "cancelling":
                raise asyncio.CancelledError
            set_status(run, "running")
            sink_id = logger.add(
                log_sink,
                level="INFO",
                format="{time:HH:mm:ss} | {level:<8} | {message}",
            )
            publish(run, "user", content=run.prompt)
            publish(run, "mode", mode=run.mode)
            if run.parent_run_id:
                publish(run, "context", parent_run_id=run.parent_run_id)
            if run.auto_skill_matches:
                publish(
                    run,
                    "skill_match",
                    count=len(run.auto_skill_matches),
                    matches=run.auto_skill_matches,
                )
            elif list_skill_documents():
                publish(
                    run,
                    "skill_match",
                    count=0,
                    matches=[],
                    message="关键词匹配未找到超过阈值的 Skill。",
                )
            if run.memory_matches:
                publish(
                    run,
                    "memory_rag",
                    count=len(run.memory_matches),
                    matches=run.memory_matches,
                )
            elif config.memory_config.enabled:
                publish(
                    run,
                    "memory_rag",
                    count=0,
                    matches=[],
                    message="Memory RAG 未召回到相关历史。",
                )

            execution_prompt = build_execution_prompt(run)
            if run.mode == "team":
                outcome = await execute_team_run(run, execution_prompt)
                usable_results = sum(
                    result.get("status") in {"completed", "partial"}
                    for result in (run.team or {}).get("results", [])
                )
                if run.status == "cancelling":
                    set_status(run, "cancelled")
                elif usable_results:
                    set_status(
                        run,
                        "completed",
                        message=(
                            "All team tasks completed."
                            if outcome.success
                            else "Team run completed with partial results."
                        ),
                    )
                else:
                    run.error = "No team worker completed its assigned task."
                    set_status(run, "error", message=run.error)
                return

            agent = await Manus.create()
            agent.max_steps = run.max_steps or DEFAULT_MAX_STEPS

            tool_names = sorted(agent.available_tools.tool_map)
            publish(
                run,
                "tools",
                count=len(tool_names),
                browser_tools=len(
                    [name for name in tool_names if name.startswith("mcp_playwright_")]
                ),
                gmail_tools=len(
                    [name for name in tool_names if name.startswith("mcp_gmail_")]
                ),
            )

            if image_records_for_direct_input(run.attachments):
                agent.memory.add_message(Message.user_message(execution_prompt))
                add_uploaded_images_to_agent(agent, run.attachments)
                result = await agent.run(task_objective=run.prompt)
            else:
                result = await agent.run(
                    execution_prompt,
                    task_objective=run.prompt,
                )
            run.result = redact(result)
            hit_step_limit = "Terminated: Reached max steps" in result
            final_answer = extract_final_answer(agent, result)
            if final_answer == NO_FINAL_ANSWER_MESSAGE:
                final_answer = result
            final_answer = redact(final_answer)
            if hit_step_limit:
                run.answer = "任务达到最大步数上限，尚未生成最终答案。请提高步数后重新运行，或基于当前页面继续任务。"
            else:
                run.answer = final_answer
            publish(run, "answer", content=run.answer)
            publish(run, "result", content=run.result)
            if run.status == "cancelling":
                set_status(run, "cancelled")
            elif hit_step_limit:
                set_status(
                    run,
                    "step_limit",
                    message="Reached max steps before producing a final answer.",
                )
            elif getattr(agent, "finish_status", None) == "failure":
                run.error = redact(
                    getattr(agent, "finish_reason", None)
                    or "Agent terminated without completing the requested task."
                )
                set_status(run, "error", message=run.error)
            elif getattr(agent, "finish_status", None) == "success":
                set_status(run, "completed")
            else:
                run.error = "Agent ended without a valid success or failure status."
                set_status(run, "error", message=run.error)
    except asyncio.CancelledError as exc:
        externally_cancelled = run.status == "cancelling"
        if externally_cancelled:
            run.error = "Run cancelled."
            set_status(run, "cancelled")
        else:
            run.error = redact(
                f"Internal cancellation occurred during execution: {str(exc) or type(exc).__name__}"
            )
            publish(run, "error", message=run.error)
            set_status(run, "error")
    except Exception as exc:
        run.error = redact(str(exc))
        publish(run, "error", message=run.error)
        set_status(run, "error")
    finally:
        if run.status in {"completed", "step_limit"}:
            try:
                await asyncio.to_thread(remember_run, run)
            except Exception as exc:
                publish(
                    run,
                    "log",
                    level="WARNING",
                    message=f"Persisting Agent Memory failed: {exc}",
                )
        if agent is not None:
            try:
                await agent.cleanup()
            except Exception as exc:
                publish(run, "log", level="WARNING", message=f"Cleanup failed: {exc}")
        if sink_id is not None:
            logger.remove(sink_id)


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/favicon.ico", include_in_schema=False)
async def favicon() -> FileResponse:
    return FileResponse(ASSETS_DIR / "logo.jpg")


@app.get("/api/files/{file_path:path}")
async def workspace_file(file_path: str) -> FileResponse:
    root = config.workspace_root.resolve()
    target = (root / file_path).resolve()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail="File not found.") from exc

    if not target.is_file():
        raise HTTPException(status_code=404, detail="File not found.")

    return FileResponse(target, filename=target.name)


@app.post("/api/uploads")
async def upload_file(file: UploadFile = File(...)) -> dict[str, Any]:
    return await save_uploaded_file(file)


def normalized_skill_content(payload: SkillRequest) -> str:
    content = payload.content.strip()
    if not re.search(r"(?m)^#\s+", content):
        content = f"# {payload.name.strip()}\n\n{content}"
    return content


@app.get("/api/skills")
async def list_skills() -> dict[str, Any]:
    return {
        "skills_dir": str(SKILLS_DIR),
        "skills": list_skill_records(),
    }


@app.post("/api/skills")
async def create_skill(payload: SkillRequest) -> dict[str, Any]:
    base_id = normalize_skill_id(payload.name)
    skill_id = base_id
    suffix = 2
    while skill_markdown_path(skill_id).exists():
        skill_id = normalize_skill_id(f"{base_id}-{suffix}")
        suffix += 1

    path = skill_markdown_path(skill_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(normalized_skill_content(payload), encoding="utf-8")
    content = path.read_text(encoding="utf-8")
    return parse_skill_metadata(skill_id, content)


@app.get("/api/skills/{skill_id}")
async def get_skill(skill_id: str) -> dict[str, Any]:
    normalized = normalize_skill_id(skill_id)
    content = read_skill_content(normalized)
    metadata = parse_skill_metadata(normalized, content)
    metadata["content"] = content
    return metadata


@app.put("/api/skills/{skill_id}")
async def update_skill(skill_id: str, payload: SkillRequest) -> dict[str, Any]:
    normalized = normalize_skill_id(skill_id)
    path = skill_markdown_path(normalized)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Skill not found.")
    path.write_text(normalized_skill_content(payload), encoding="utf-8")
    content = path.read_text(encoding="utf-8")
    metadata = parse_skill_metadata(normalized, content)
    metadata["content"] = content
    return metadata


@app.delete("/api/skills/{skill_id}")
async def delete_skill(skill_id: str) -> dict[str, Any]:
    normalized = normalize_skill_id(skill_id)
    path = skill_markdown_path(normalized)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Skill not found.")
    path.unlink()
    try:
        path.parent.rmdir()
    except OSError:
        pass
    return {"deleted": normalized}


@app.get("/api/status")
async def status() -> dict[str, Any]:
    mcp_servers = {}
    for name, server in config.mcp_config.servers.items():
        mcp_servers[name] = {
            "type": server.type,
            "command": server.command,
            "url": server.url,
            "args": server.args,
            "env": {key: "<redacted>" for key in (server.env or {}).keys()},
            "headers": {key: "<redacted>" for key in (server.headers or {}).keys()},
        }

    default_llm = config.llm.get("default")
    return {
        "workspace": str(config.workspace_root),
        "project_root": str(PROJECT_ROOT),
        "model": getattr(default_llm, "model", None),
        "base_url": getattr(default_llm, "base_url", None),
        "reasoning_effort": getattr(default_llm, "reasoning_effort", None),
        "execution_modes": ["single", "team"],
        "mcp_servers": mcp_servers,
        "gmail": gmail_auth_status(),
        "skills": list_skill_records(),
        "active_runs": [
            run.to_dict()
            for run in sorted(
                runs.values(), key=lambda item: item.updated_at, reverse=True
            )
            if run.status not in TERMINAL_STATUSES
        ],
        "runs": [run.to_dict() for run in recent_thread_runs()],
    }


@app.get("/api/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/runs")
async def create_run(payload: RunRequest) -> dict[str, Any]:
    prompt = payload.prompt.strip()
    if not prompt:
        raise HTTPException(status_code=400, detail="Prompt cannot be empty.")

    parent_run_id = payload.parent_run_id.strip() if payload.parent_run_id else None
    if parent_run_id:
        parent = runs.get(parent_run_id)
        if not parent:
            raise HTTPException(status_code=404, detail="Parent run not found.")
        if parent.status not in TERMINAL_STATUSES:
            raise HTTPException(status_code=409, detail="Parent run is still active.")
    attachment_ids = list(dict.fromkeys(payload.attachment_ids))
    if len(attachment_ids) > MAX_UPLOADS_PER_RUN:
        raise HTTPException(
            status_code=400,
            detail=f"At most {MAX_UPLOADS_PER_RUN} attachments per run.",
        )
    attachments = selected_upload_records(attachment_ids)
    if len(attachments) != len(attachment_ids):
        raise HTTPException(
            status_code=400, detail="One or more uploaded attachments were not found."
        )

    run_id = uuid.uuid4().hex[:12]
    now = utc_now()
    run = RunSession(
        id=run_id,
        prompt=prompt,
        max_steps=DEFAULT_MAX_STEPS,
        parent_run_id=parent_run_id,
        skill_ids=valid_skill_ids(payload.skill_ids),
        auto_skill_matches=[],
        attachments=attachments,
        mode=payload.mode,
        created_at=now,
        updated_at=now,
    )
    run.auto_skill_matches = await retrieve_auto_skill_matches(
        prompt, run.skill_ids, attachments
    )
    recent_history = conversation_runs(parent_run_id)[-8:] if parent_run_id else []
    run.memory_matches = await retrieve_relevant_memories(
        prompt,
        exclude_run_ids={item.id for item in recent_history},
    )
    runs[run_id] = run
    prune_runs()
    publish(run, "created", id=run_id, mode=run.mode)
    run.task = asyncio.create_task(execute_run(run))
    return run.to_dict()


@app.get("/api/runs/{run_id}")
async def get_run(run_id: str) -> dict[str, Any]:
    run = runs.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found.")
    return {
        **run.to_dict(),
        "events": run.events,
        "conversation": [item.to_dict() for item in conversation_runs(run_id)],
    }


@app.post("/api/runs/{run_id}/cancel")
async def cancel_run(run_id: str) -> dict[str, Any]:
    run = runs.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found.")
    if run.status in TERMINAL_STATUSES:
        return run.to_dict()
    if run.status != "cancelling":
        set_status(run, "cancelling", message="Stopping the current run...")
    if run.task:
        run.task.cancel()
    else:
        set_status(run, "cancelled")
    return run.to_dict()


@app.delete("/api/runs")
async def delete_all_finished_runs() -> dict[str, Any]:
    delete_ids: set[str] = set()
    skipped_active: list[str] = []

    for recent_run in recent_thread_runs():
        thread_ids = thread_run_ids(recent_run.id)
        if any(
            runs[thread_id].status not in TERMINAL_STATUSES for thread_id in thread_ids
        ):
            skipped_active.append(recent_run.id)
            continue
        delete_ids.update(thread_ids)

    for run_id in delete_ids:
        runs.pop(run_id, None)
    if config.memory_config.enabled:
        await asyncio.to_thread(MEMORY_STORE.delete_run_ids, delete_ids)
    save_runs()
    return {
        "deleted": sorted(delete_ids),
        "skipped_active": skipped_active,
    }


@app.delete("/api/runs/{run_id}")
async def delete_run(run_id: str) -> dict[str, Any]:
    run = runs.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found.")

    delete_ids = thread_run_ids(run_id)
    active_ids = [
        candidate_id
        for candidate_id in delete_ids
        if runs[candidate_id].status not in TERMINAL_STATUSES
    ]
    if active_ids:
        raise HTTPException(
            status_code=409, detail="Cannot delete an active conversation."
        )

    for candidate_id in delete_ids:
        runs.pop(candidate_id, None)
    if config.memory_config.enabled:
        await asyncio.to_thread(MEMORY_STORE.delete_run_ids, set(delete_ids))
    save_runs()
    return {"deleted": delete_ids}


@app.websocket("/ws/runs/{run_id}")
async def run_events(websocket: WebSocket, run_id: str) -> None:
    run = runs.get(run_id)
    if not run:
        await websocket.close(code=4404)
        return

    await websocket.accept()
    queue: asyncio.Queue = asyncio.Queue(maxsize=200)
    run.subscribers.add(queue)
    try:
        for event in run.events:
            await websocket.send_json(event)
        while True:
            event = await queue.get()
            await websocket.send_json(event)
            if event["type"] == "status" and event.get("status") in TERMINAL_STATUSES:
                break
    except WebSocketDisconnect:
        pass
    finally:
        run.subscribers.discard(queue)


def run_web(host: str = "127.0.0.1", port: int = 7788, reload: bool = False) -> None:
    import uvicorn

    os.environ["OPENMANUS_WEB_MODE"] = "1"
    uvicorn.run(
        "app.web.server:app",
        host=host,
        port=port,
        reload=reload,
        log_level="info",
    )
