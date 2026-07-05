import json
import re
from pathlib import Path
from typing import Any, Optional

from app.config import config
from app.exceptions import ToolError
from app.tool.base import BaseTool, ToolResult


def _resolve_workspace_file(output_path: Optional[str], default_name: str, suffix: str) -> Path:
    raw_path = output_path or default_name
    path = Path(raw_path).expanduser()
    if not path.is_absolute():
        path = config.workspace_root / path

    if path.suffix.lower() != suffix:
        path = path.with_suffix(suffix)

    resolved = path.resolve()
    workspace = config.workspace_root.resolve()
    try:
        resolved.relative_to(workspace)
    except ValueError as exc:
        raise ToolError(
            f"Output path must be inside the workspace: {workspace}"
        ) from exc

    resolved.parent.mkdir(parents=True, exist_ok=True)
    return resolved


def _json_result(data: dict[str, Any]) -> ToolResult:
    return ToolResult(output=json.dumps(data, ensure_ascii=False, indent=2))


class CreateWordDocument(BaseTool):
    name: str = "create_word_document"
    description: str = (
        "Create a Microsoft Word .docx document in the workspace. "
        "Use this when the user asks for a report, proposal, memo, resume, paper, "
        "meeting notes, or any deliverable that should be opened and edited in Word. "
        "Supports headings, paragraphs, bullet lists, numbered lists, tables, and page breaks."
    )
    parameters: dict = {
        "type": "object",
        "properties": {
            "output_path": {
                "type": "string",
                "description": "Workspace-relative or absolute output path. The .docx suffix is added if omitted.",
            },
            "title": {
                "type": "string",
                "description": "Optional document title. Added as the first Word title paragraph.",
            },
            "author": {
                "type": "string",
                "description": "Optional document author metadata.",
            },
            "content": {
                "type": "string",
                "description": "Optional plain text content. Used when sections are not provided; blank lines become paragraph breaks.",
            },
            "sections": {
                "type": "array",
                "description": "Structured document blocks in order.",
                "items": {
                    "type": "object",
                    "properties": {
                        "type": {
                            "type": "string",
                            "enum": [
                                "heading",
                                "paragraph",
                                "bullets",
                                "numbered_list",
                                "table",
                                "page_break",
                            ],
                        },
                        "text": {"type": "string"},
                        "level": {
                            "type": "integer",
                            "description": "Heading level from 1 to 9.",
                        },
                        "items": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "headers": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "rows": {
                            "type": "array",
                            "items": {
                                "type": "array",
                                "items": {
                                    "type": ["string", "number", "integer", "boolean", "null"]
                                },
                            },
                        },
                    },
                    "required": ["type"],
                },
            },
            "overwrite": {
                "type": "boolean",
                "description": "Whether to overwrite an existing file. Defaults to true.",
            },
        },
        "required": [],
    }

    async def execute(
        self,
        output_path: Optional[str] = None,
        title: Optional[str] = None,
        author: Optional[str] = None,
        content: Optional[str] = None,
        sections: Optional[list[dict[str, Any]]] = None,
        overwrite: bool = True,
        **_: Any,
    ) -> ToolResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise ToolError("python-docx is not installed. Install python-docx first.") from exc

        path = _resolve_workspace_file(output_path, "document.docx", ".docx")
        if path.exists() and not overwrite:
            raise ToolError(f"File already exists: {path}")

        document = Document()
        if title:
            document.core_properties.title = title
            document.add_heading(title, level=0)
        if author:
            document.core_properties.author = author

        block_count = 0
        if sections:
            for section in sections:
                block_count += 1
                self._add_section(document, section)
        elif content:
            block_count = self._add_plain_content(document, content)
        elif not title:
            document.add_paragraph("")

        document.save(path)
        return _json_result(
            {
                "ok": True,
                "type": "word_document",
                "path": str(path),
                "blocks": block_count,
                "message": "Word document created successfully.",
            }
        )

    def _add_section(self, document: Any, section: dict[str, Any]) -> None:
        section_type = section.get("type")
        if section_type == "heading":
            level = max(1, min(9, int(section.get("level") or 1)))
            document.add_heading(str(section.get("text") or ""), level=level)
        elif section_type == "paragraph":
            document.add_paragraph(str(section.get("text") or ""))
        elif section_type == "bullets":
            for item in section.get("items") or []:
                document.add_paragraph(str(item), style="List Bullet")
        elif section_type == "numbered_list":
            for item in section.get("items") or []:
                document.add_paragraph(str(item), style="List Number")
        elif section_type == "table":
            self._add_table(document, section)
        elif section_type == "page_break":
            document.add_page_break()
        else:
            raise ToolError(f"Unsupported Word section type: {section_type}")

    def _add_plain_content(self, document: Any, content: str) -> int:
        blocks = 0
        for raw_block in re.split(r"\n\s*\n", content.strip()):
            text = raw_block.strip()
            if not text:
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", text)
            if heading:
                document.add_heading(heading.group(2), level=len(heading.group(1)))
            else:
                document.add_paragraph(text)
            blocks += 1
        return blocks

    def _add_table(self, document: Any, section: dict[str, Any]) -> None:
        headers = [str(item) for item in (section.get("headers") or [])]
        rows = section.get("rows") or []
        if not headers and not rows:
            raise ToolError("Word table section requires headers or rows.")

        column_count = len(headers) or max(len(row) for row in rows)
        table = document.add_table(rows=1 if headers else 0, cols=column_count)
        table.style = "Table Grid"

        if headers:
            header_cells = table.rows[0].cells
            for index in range(column_count):
                header_cells[index].text = headers[index] if index < len(headers) else ""

        for row in rows:
            cells = table.add_row().cells
            for index in range(column_count):
                value = row[index] if index < len(row) else ""
                cells[index].text = "" if value is None else str(value)


class CreateExcelWorkbook(BaseTool):
    name: str = "create_excel_workbook"
    description: str = (
        "Create a Microsoft Excel .xlsx workbook in the workspace. "
        "Use this when the user asks for spreadsheets, tables, data workbooks, trackers, "
        "budgets, schedules, exports, or any deliverable that should be opened and edited in Excel. "
        "Supports multiple sheets, headers, rows, formulas, auto filters, frozen panes, and column widths."
    )
    parameters: dict = {
        "type": "object",
        "properties": {
            "output_path": {
                "type": "string",
                "description": "Workspace-relative or absolute output path. The .xlsx suffix is added if omitted.",
            },
            "sheets": {
                "type": "array",
                "description": "Workbook sheets in order.",
                "items": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string"},
                        "headers": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "rows": {
                            "type": "array",
                            "items": {"type": ["array", "object"]},
                        },
                        "freeze_panes": {
                            "type": "string",
                            "description": "Excel cell reference such as A2. Use an empty string to disable.",
                        },
                        "auto_filter": {"type": "boolean"},
                        "table": {
                            "type": "boolean",
                            "description": "Whether to format the range as an Excel table when headers are present.",
                        },
                        "column_widths": {
                            "type": "object",
                            "description": "Column widths keyed by Excel letter (A) or header name.",
                        },
                    },
                    "required": ["name"],
                },
            },
            "overwrite": {
                "type": "boolean",
                "description": "Whether to overwrite an existing file. Defaults to true.",
            },
        },
        "required": ["sheets"],
    }

    async def execute(
        self,
        sheets: list[dict[str, Any]],
        output_path: Optional[str] = None,
        overwrite: bool = True,
        **_: Any,
    ) -> ToolResult:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            from openpyxl.utils import get_column_letter
            from openpyxl.worksheet.table import Table, TableStyleInfo
        except ImportError as exc:
            raise ToolError("openpyxl is not installed. Install openpyxl first.") from exc

        if not sheets:
            raise ToolError("At least one sheet is required.")

        path = _resolve_workspace_file(output_path, "workbook.xlsx", ".xlsx")
        if path.exists() and not overwrite:
            raise ToolError(f"File already exists: {path}")

        workbook = Workbook()
        workbook.remove(workbook.active)

        total_rows = 0
        for index, sheet in enumerate(sheets, start=1):
            name = self._safe_sheet_name(sheet.get("name") or f"Sheet{index}")
            worksheet = workbook.create_sheet(name)
            headers, rows = self._normalize_rows(sheet.get("headers"), sheet.get("rows") or [])

            if headers:
                worksheet.append(headers)
                for cell in worksheet[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="1F5F99")

            for row in rows:
                worksheet.append(row)

            row_count = len(rows) + (1 if headers else 0)
            col_count = max((len(row) for row in ([headers] if headers else []) + rows), default=1)
            total_rows += len(rows)

            if row_count and col_count:
                worksheet.auto_filter.ref = (
                    f"A1:{get_column_letter(col_count)}{row_count}"
                    if sheet.get("auto_filter", bool(headers))
                    else None
                )

            freeze_panes = sheet.get("freeze_panes")
            if freeze_panes is None and headers:
                freeze_panes = "A2"
            if freeze_panes:
                worksheet.freeze_panes = freeze_panes

            if headers and rows and sheet.get("table", True):
                ref = f"A1:{get_column_letter(col_count)}{row_count}"
                table = Table(displayName=f"Table{index}", ref=ref)
                table.tableStyleInfo = TableStyleInfo(
                    name="TableStyleMedium2",
                    showFirstColumn=False,
                    showLastColumn=False,
                    showRowStripes=True,
                    showColumnStripes=False,
                )
                worksheet.add_table(table)

            self._apply_widths(worksheet, headers, rows, sheet.get("column_widths") or {})

        workbook.save(path)
        return _json_result(
            {
                "ok": True,
                "type": "excel_workbook",
                "path": str(path),
                "sheets": [sheet.get("name") for sheet in sheets],
                "data_rows": total_rows,
                "message": "Excel workbook created successfully.",
            }
        )

    def _normalize_rows(
        self, headers: Optional[list[str]], rows: list[Any]
    ) -> tuple[list[str], list[list[Any]]]:
        header_values = [str(item) for item in headers] if headers else []
        if not header_values:
            seen: list[str] = []
            for row in rows:
                if isinstance(row, dict):
                    for key in row.keys():
                        key_text = str(key)
                        if key_text not in seen:
                            seen.append(key_text)
            header_values = seen

        normalized: list[list[Any]] = []
        for row in rows:
            if isinstance(row, dict):
                normalized.append([row.get(header) for header in header_values])
            elif isinstance(row, list):
                normalized.append(row)
            else:
                raise ToolError("Excel rows must be arrays or objects.")

        return header_values, normalized

    def _safe_sheet_name(self, value: str) -> str:
        name = re.sub(r"[:\\/?*\[\]]", "_", str(value)).strip() or "Sheet"
        return name[:31]

    def _apply_widths(
        self,
        worksheet: Any,
        headers: list[str],
        rows: list[list[Any]],
        explicit_widths: dict[str, Any],
    ) -> None:
        from openpyxl.utils import get_column_letter

        all_rows = ([headers] if headers else []) + rows
        max_cols = max((len(row) for row in all_rows), default=0)
        header_to_letter = {
            header: get_column_letter(index + 1) for index, header in enumerate(headers)
        }

        for col_index in range(1, max_cols + 1):
            letter = get_column_letter(col_index)
            width = explicit_widths.get(letter)
            if width is None and col_index <= len(headers):
                width = explicit_widths.get(headers[col_index - 1])

            if width is None:
                sample_values = [
                    "" if col_index > len(row) or row[col_index - 1] is None else str(row[col_index - 1])
                    for row in all_rows
                ]
                width = min(max([len(value) for value in sample_values] + [10]) + 2, 60)

            worksheet.column_dimensions[letter].width = float(width)

        for header, letter in header_to_letter.items():
            if header in explicit_widths:
                worksheet.column_dimensions[letter].width = float(explicit_widths[header])
